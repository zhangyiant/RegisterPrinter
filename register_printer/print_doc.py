import re
import os
import os.path
import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


LOGGER = logging.getLogger(__name__)


def print_doc_reg(reg, dh, reg_idx, blk_idx, blk_insts):
    dh.add_page_break()
    dh.add_heading("  %d.%d  %s" % (blk_idx, reg_idx + 1, reg.name), level=2)
    p = dh.add_paragraph()
    p.add_run('    Offset : ').bold = True
    p.add_run('%s\n' % (hex(reg.offset)))

    for blk_inst in blk_insts:
        p.add_run('    %s Address : ' % (blk_inst['inst_name'])).bold = True
        p.add_run('%s\n' % (hex(blk_inst['inst_base'] + reg.offset)))

    p.add_run("    Reset Value : ").bold = True
    p.add_run("0x%x\n" % (reg.calculate_register_default()))

    p.add_run("    Description : ").bold = True
    p.add_run("%s\n" % (reg.description))

    headers = [
        "LSB",
        "MSB",
        "Field",
        "Access",
        "Default",
        "Description"]
    tb = dh.add_table(
        len(reg.fields) + 1,
        len(headers),
        style="Light Grid")
    tb.autofit = 1
    hcells = tb.rows[0].cells
    i = 0
    for header in headers:
        p = hcells[i].paragraphs[0]
        run = p.add_run(header)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1

    i = 1
    for field in reg.fields:
        tb.cell(i, 0).text = str(field.lsb)
        tb.cell(i, 1).text = str(field.msb)
        tb.cell(i, 2).text = str(field.name)
        tb.cell(i, 3).text = str(field.access)
        tb.cell(i, 4).text = hex(field.default)
        tb.cell(i, 5).text = "%s" % (field.description)
        i += 1
    return

def print_doc_block(block, doc, idx, instances):
    doc.add_heading("%d %s Registers" % (idx, block.name), level=1)
    tb = doc.add_table(
        len(block.registers) + 1,
        2 + len(instances),
        style="Light Grid")
    hcells = tb.rows[0].cells
    hdr = []
    hdr.append("Offset")
    for instance in instances:
        hdr.append("%s Addr" % (instance["inst_name"]))
    hdr.append("Register")
    for i in range(len(hdr)):
        p = hcells[i].paragraphs[0]
        run = p.add_run(hdr[i])
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    i = 1
    for register in block.registers:
        tb.cell(i, 0).text = hex(register.offset)
        for k in range(len(instances)):
            tb.cell(i, k+1).text = hex(
                instances[k]['inst_base'] + register.offset)
        tb.cell(i, len(hdr) - 1).text = str(register.name)
        i += 1

    reg_idx = 0
    for register in block.registers:
        print_doc_reg(register, doc, reg_idx, idx, instances)
        reg_idx = reg_idx + 1
    doc.add_page_break()
    return

def generate_doc(top_sys):
    doc = Document()

    title = doc.add_heading(
        top_sys.name + " Registers",
        level=0)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Version : %s\n" % (top_sys.version))
    p.add_run("Author : %s" % (top_sys.author))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    block_idx = 1
    doc.add_heading("%d Address Map" % block_idx, level=1)
    table = doc.add_table(
        len(top_sys.addr_map) + 1,
        3,
        style="Light Grid")
    hcell = table.rows[0].cells
    headers = ["Block", "Start Address", "Size"]
    for i in range(3):
        p = hcell[i].paragraphs[0]
        run = p.add_run(headers[i])
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    i = 1
    for addr_map in top_sys.addr_map:
        table.cell(i, 0).text = str(addr_map['block_instance'])
        table.cell(i, 1).text = str(addr_map['base_address'])
        table.cell(i, 2).text = hex(addr_map['block_size'])
        i += 1
    doc.add_page_break()

    block_idx += 1
    for block in top_sys.blocks:
        # get block instances
        blk_insts = []
        for addr_entry in top_sys.addr_map:
            if addr_entry['block_type'] == block.name:
                inst = {
                    "inst_name": addr_entry["block_instance"],
                    "inst_base": int(addr_entry["base_address"])
                }
                blk_insts.append(inst)

        print_doc_block(block, doc, block_idx, blk_insts)
        block_idx = block_idx + 1
    return doc

def print_doc(top_sys, output_path="."):
    LOGGER.debug("Generating register description document...")

    doc_file_name = os.path.join(
        output_path,
        top_sys.name.lower() + "_registers.docx"
    )
    if os.path.exists(doc_file_name):
        os.remove(doc_file_name)

    doc = generate_doc(top_sys)

    doc.save(doc_file_name)

    LOGGER.debug("Register description document is generated %s", doc_file_name)
    return
