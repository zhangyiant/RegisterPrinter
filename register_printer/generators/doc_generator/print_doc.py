import os
import os.path
import logging
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from register_printer.data_model import Register, Array, Field

LOGGER = logging.getLogger(__name__)


def print_doc_reg(doc, blk_idx, reg_idx, register, block_instances,
                  array=None):
    doc.add_page_break()
    doc.add_heading(
        "  %d.%d  %s" % (blk_idx, reg_idx + 1, register.name),
        level=2
    )
    p = doc.add_paragraph()

    add_register_offset(p, register, array)

    add_register_instances_address(p, block_instances, register, array)

    add_register_reset_value(p, register, array)

    p.add_run("    Description : ").bold = True
    p.add_run("%s\n" % (register.description))

    headers = [
        "LSB",
        "MSB",
        "Field",
        "Access",
        "Default",
        "Description"]
    row_count = 1
    for field in register.fields:
        if field.name != "-":
            row_count += 1
    tb = doc.add_table(
        row_count,
        len(headers),
        style="Light Grid")
    tb.autofit = 1
    hcells = tb.rows[0].cells
    i = 0
    for header in headers:
        p = hcells[i].paragraphs[0]
        p.add_run(header)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        i += 1

    i = 1
    for field in register.fields:
        if field.name == "-":
            continue
        tb.cell(i, 0).text = str(field.lsb)
        tb.cell(i, 1).text = str(field.msb)
        tb.cell(i, 2).text = str(field.name)
        tb.cell(i, 3).text = str(field.access)
        tb.cell(i, 4).text = hex(field.default)
        tb.cell(i, 5).text = "%s" % (field.description)
        i += 1
    return


def add_register_reset_value(paragraph, register, array):
    paragraph.add_run("    Reset Value : ").bold = True
    paragraph.add_run("0x%x\n" % (register.calculate_register_default()))
    if array is not None:
        for index in range(array.length):
            offset = array.start_address + array.offset * index + \
                     register.offset
            temp_reg = Register(offset)
            temp_reg.name = register.name
            for field in register.fields:
                temp_reg.fields.append(field)
            found = False
            for overwrite_entry in array.default_overwrite_entries:
                if overwrite_entry.index == index and \
                        overwrite_entry.register_name == temp_reg.name:
                    found = True
                    field_name = overwrite_entry.field_name
                    for field in temp_reg.fields:
                        if field.name == field_name:
                            temp_reg.fields.remove(field)
                            new_field = Field()
                            new_field.name = field.name
                            new_field.msb = field.msb
                            new_field.lsb = field.lsb
                            new_field.default = overwrite_entry.default
                            new_field.access = field.access
                            new_field.description = field.description
                            temp_reg.fields.append(new_field)
            if found:
                paragraph.add_run(
                    "           " +
                    hex(temp_reg.offset) + ":0x%x\n" %
                    temp_reg.calculate_register_default()
                )


def add_register_instances_address(paragraph, block_instances, register,
                                   array):
    for blk_inst in block_instances:
        paragraph.add_run('    %s Address : ' % (blk_inst.name)).bold = True
        if array is None:
            paragraph.add_run(
                '%s\n' % (hex(blk_inst.base_address + register.offset))
            )
        else:
            addr = blk_inst.base_address + array.start_address + \
                   register.offset
            paragraph.add_run('%s\n' % hex(addr))


def add_register_offset(paragraph, register, array):
    paragraph.add_run('    Offset : ').bold = True
    if array is None:
        paragraph.add_run('%s\n' % (hex(register.offset)))
    else:
        offset_str = hex(
            register.offset + array.start_address)
        offset_str += " + " + hex(array.offset) + " * n"
        offset_str += " (n >= 0, n < " + str(array.length) + ")"
        paragraph.add_run('%s\n' % offset_str)


def get_registers(block):
    result = []
    for register in block.registers:
        if isinstance(register, Array):
            result.append(register)
        elif isinstance(register, Register):
            if register.is_reserved:
                continue
            else:
                result.append(register)
    return result


def calc_row_num(registers):
    num = 0
    for register in registers:
        if isinstance(register, Array):
            struct = register.content_type
            for struct_register in struct.registers:
                if struct_register.is_reserved:
                    continue
                else:
                    num += 1
        else:
            num += 1
    return num


def print_doc_block(doc, idx, block, instances):
    block_type = block.block_type

    registers = get_registers(block)

    doc.add_heading("%d %s Registers" % (idx, block_type), level=1)

    num_register = calc_row_num(registers)
    row_count = num_register + 1
    column_count = 2 + len(instances)
    tb = doc.add_table(
        row_count,
        column_count,
        style="Light Grid")

    hdr = get_table_header(instances)

    hcells = tb.rows[0].cells
    for i in range(column_count):
        p = hcells[i].paragraphs[0]
        p.add_run(hdr[i])
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    i = 1
    for register in registers:
        if isinstance(register, Register):
            tb.cell(i, 0).text = hex(register.offset)
            for k in range(len(instances)):
                tb.cell(i, k + 1).text = hex(
                    instances[k].base_address + register.offset)
            tb.cell(i, column_count - 1).text = str(register.name)
            i += 1
        elif isinstance(register, Array):
            struct = register.content_type
            for struct_registers in struct.registers:
                if struct_registers.is_reserved:
                    continue
                offset_str = hex(
                    struct_registers.offset + register.start_address)
                offset_str += " + " + hex(register.offset) + " * n"
                offset_str += " (n >= 0, n < " + str(register.length) + ")"
                tb.cell(i, 0).text = offset_str
                for k in range(len(instances)):
                    tb.cell(i, k + 1).text = hex(
                        instances[k].base_address + register.start_address +
                        struct_registers.offset
                    )
                tb.cell(i, column_count - 1).text = str(struct_registers.name)
                i += 1

    add_block_registers(doc, idx, registers, instances)

    doc.add_page_break()
    return


def get_table_header(instances):
    hdr = ["Offset"]
    for instance in instances:
        hdr.append("%s Addr" % instance.name)
    hdr.append("Register")
    return hdr


def add_block_registers(doc, idx, registers, instances):
    reg_idx = 0
    for register in registers:
        if isinstance(register, Register):
            print_doc_reg(doc, idx, reg_idx, register, instances)
            reg_idx += 1
        if isinstance(register, Array):
            struct = register.content_type
            for struct_reg in struct.registers:
                if struct_reg.is_reserved:
                    continue
                print_doc_reg(doc, idx, reg_idx, struct_reg, instances,
                              register)
                reg_idx += 1
    return


def generate_doc(top_sys):
    doc = Document()

    add_header(doc, top_sys.name)

    add_version_author(doc, top_sys.version, top_sys.author)

    doc.add_page_break()

    block_idx = 1
    block_instances = top_sys.block_instances
    add_address_map(doc, block_idx, block_instances)

    doc.add_page_break()

    block_idx += 1
    for block in top_sys.blocks:
        block_type = block.block_type
        blk_insts = []
        for instance in top_sys.block_instances:
            if instance.block_type == block_type:
                blk_insts.append(instance)

        print_doc_block(doc, block_idx, block, blk_insts)
        block_idx = block_idx + 1
    return doc


def add_address_map(doc, block_idx, block_instances):
    doc.add_heading("%d Address Map" % block_idx, level=1)
    table = doc.add_table(
        len(block_instances) + 1,
        3,
        style="Light Grid")
    hcell = table.rows[0].cells
    headers = ["Block", "Start Address", "Size"]
    for i in range(3):
        p = hcell[i].paragraphs[0]
        p.add_run(headers[i])
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    i = 1
    for block_instance in block_instances:
        table.cell(i, 0).text = block_instance.name
        table.cell(i, 1).text = hex(block_instance.base_address)
        table.cell(i, 2).text = hex(block_instance.size)
        i += 1


def add_version_author(doc, version, author):
    p = doc.add_paragraph()
    p.add_run("Version : %s\n" % version)
    p.add_run("Author : %s" % author)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_header(doc, name):
    title = doc.add_heading(
        name + " Registers",
        level=0)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return


def print_doc(top_sys, output_path="."):
    LOGGER.debug("Generating register description document...")

    doc_file_name = os.path.join(
        output_path,
        top_sys.name.lower() + "_registers.docx"
    )
    if os.path.exists(doc_file_name):
        os.remove(doc_file_name)

    doc = generate_doc(top_sys)

    doc.save(doc_file_name)

    LOGGER.debug(
        "Register description document is generated %s",
        doc_file_name
    )
    return
